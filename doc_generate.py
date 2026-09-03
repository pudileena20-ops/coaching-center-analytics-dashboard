from docx import Document

doc = Document()

# Title
doc.add_heading('AI Hospital Receptionist System', level=1)
doc.add_paragraph('High-Level Architecture and Flow Documentation')

# Overview
doc.add_heading('1. Project Overview', level=2)
doc.add_paragraph(
    "This system is a multilingual AI-powered hospital receptionist built using "
    "FastAPI, Google Gemini (LLM), SQLite, and gTTS for audio generation. "
    "It allows users to chat with an AI agent to list doctors, check availability, "
    "book appointments, and receive responses in text and audio format."
)

# Components
doc.add_heading('2. High-Level Components', level=2)

components = [
    "Frontend (HTML UI) – Collects user input and plays audio responses.",
    "FastAPI Backend – Handles API routes and session management.",
    "Gemini LLM – Processes user messages and performs tool-calling.",
    "Tool Layer – Python functions for doctor listing, slot checking, and booking.",
    "SQLite Database – Stores doctors and appointments.",
    "gTTS Module – Converts AI text responses into speech (MP3, Base64 encoded)."
]

for c in components:
    doc.add_paragraph(c, style='List Bullet')

# Flow Steps
doc.add_heading('3. High-Level System Flow', level=2)

flow_steps = [
    "1. User opens the frontend in browser.",
    "2. Frontend creates a new AI session via /agent/new_session.",
    "3. User sends a message to /agent/message endpoint.",
    "4. Backend retrieves or creates Gemini chat session.",
    "5. Language instruction is appended based on selected language.",
    "6. Gemini processes the message.",
    "7. If required, Gemini triggers a tool call (list, check, or book).",
    "8. Backend executes the tool and queries SQLite database.",
    "9. Tool result is sent back to Gemini.",
    "10. Gemini generates final natural language response.",
    "11. Response text is cleaned.",
    "12. gTTS generates audio.",
    "13. Backend returns text + audio to frontend.",
    "14. Frontend displays text and plays audio."
]

for step in flow_steps:
    doc.add_paragraph(step, style='List Number')

# Flow Diagram
doc.add_heading('4. Flow Diagram', level=2)

diagram = """
User
  |
  v
Frontend (HTML UI)
  |
  v
FastAPI Backend
  |
  v
Gemini LLM  <---- Tool Calls ---->  Python Tool Functions
  |                                      |
  |                                      v
  |                                 SQLite Database
  |
  v
gTTS (Text-to-Speech)
  |
  v
Frontend (Text + Audio Response)
"""

doc.add_paragraph(diagram)

# Key Features
doc.add_heading('5. Key Features', level=2)

features = [
    "Session-based conversation memory.",
    "Multilingual enforcement (English, Hindi, Telugu).",
    "Function calling with Gemini.",
    "Database-backed appointment booking.",
    "Audio response generation using gTTS.",
    "Asynchronous processing with asyncio."
]

for f in features:
    doc.add_paragraph(f, style='List Bullet')

# Save file
doc.save("AI_Hospital_Receptionist_System_Documentation.docx")

print("Document generated successfully!")
